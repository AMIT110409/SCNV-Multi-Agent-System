import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Title Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_rows = []

    for line in lines:
        line = line.strip()
        
        # Handle Headings
        if line.startswith('# '):
            if in_table:
                process_table(doc, table_rows)
                in_table = False
                table_rows = []
            h = doc.add_heading(line[2:], level=1)
            continue
        elif line.startswith('## '):
            if in_table:
                process_table(doc, table_rows)
                in_table = False
                table_rows = []
            h = doc.add_heading(line[3:], level=2)
            continue
        elif line.startswith('### '):
            if in_table:
                process_table(doc, table_rows)
                in_table = False
                table_rows = []
            h = doc.add_heading(line[4:], level=3)
            continue
            
        # Handle Tables
        if '|' in line:
            if '---' in line:
                continue
            in_table = True
            cols = [c.strip() for c in line.split('|') if c.strip() or (line.startswith('|') and line.endswith('|'))]
            if cols:
                table_rows.append(cols)
            continue
        elif in_table and line == '':
            process_table(doc, table_rows)
            in_table = False
            table_rows = []
            continue
        
        if line == '':
            continue
            
        # Handle regular text (basic)
        if not in_table:
            doc.add_paragraph(line)

    if in_table:
        process_table(doc, table_rows)

    doc.save(docx_path)

def process_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            if j < len(table.columns):
                table.cell(i, j).text = val

if __name__ == "__main__":
    import sys
    import argparse
    
    parser = argparse.ArgumentParser(description='Convert Markdown to DOCX')
    parser.add_argument('--input', help='Input MD file path')
    parser.add_argument('--output', help='Output DOCX file path')
    
    args = parser.parse_args()
    
    if args.input and args.output:
        convert_md_to_docx(args.input, args.output)
        print(f"Successfully created {args.output}")
    else:
        # Fallback to defaults
        md_file = r"C:\Users\Abcom\Downloads\SCNV-Multi-Agent-System\SCNV-Multi-Agent-System-Anushka\docs\SCNV_SAP_Data_Guide.md"
        docx_file = r"C:\Users\Abcom\Downloads\SCNV-Multi-Agent-System\SCNV-Multi-Agent-System-Anushka\docs\SCNV_SAP_Data_Guide.docx"
        convert_md_to_docx(md_file, docx_file)
        print(f"Successfully created {docx_file}")
